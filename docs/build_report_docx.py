"""
Reads docs/report_vi.md, restructures into Vietnamese academic format, outputs docs/report_vi.docx
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree
import latex2mathml.converter
import re, os

os.chdir(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.line_spacing = 1.5

# Helper: add heading
def h(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16 if level==0 else 14 if level==1 else 13)
    return heading

# Helper: add normal paragraph with first-line indent
def p(text):
    para = doc.add_paragraph(text)
    para.paragraph_format.first_line_indent = Cm(1.27)
    para.paragraph_format.line_spacing = 1.5
    for run in para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
    return para

# Helper: add bold paragraph
def pb(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    para.paragraph_format.line_spacing = 1.5
    return para

# Helper: add code block
def add_code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    para.paragraph_format.line_spacing = 1.0
    return para

# --- Math (LaTeX -> OMML) helpers ---
def _make_omml_run(text):
    """Create an OMML run element with text"""
    r = OxmlElement('m:r')
    t = OxmlElement('m:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    return r

def _convert_mathml_node(node):
    """Recursively convert MathML node to OMML elements"""
    tag = etree.QName(node.tag).localname if '}' in node.tag else node.tag
    results = []

    if tag in ('math', 'mrow', 'mstyle', 'mpadded'):
        for child in node:
            results.extend(_convert_mathml_node(child))
    elif tag in ('mi', 'mn', 'mo'):
        text = node.text or ''
        if text:
            results.append(_make_omml_run(text))
    elif tag == 'mtext':
        text = node.text or ''
        if text:
            results.append(_make_omml_run(text))
    elif tag == 'msup':
        children = list(node)
        if len(children) >= 2:
            ssup = OxmlElement('m:sSup')
            e = OxmlElement('m:e')
            for part in _convert_mathml_node(children[0]):
                e.append(part)
            ssup.append(e)
            sup = OxmlElement('m:sup')
            for part in _convert_mathml_node(children[1]):
                sup.append(part)
            ssup.append(sup)
            results.append(ssup)
    elif tag == 'msub':
        children = list(node)
        if len(children) >= 2:
            ssub = OxmlElement('m:sSub')
            e = OxmlElement('m:e')
            for part in _convert_mathml_node(children[0]):
                e.append(part)
            ssub.append(e)
            sub = OxmlElement('m:sub')
            for part in _convert_mathml_node(children[1]):
                sub.append(part)
            ssub.append(sub)
            results.append(ssub)
    elif tag == 'msubsup':
        children = list(node)
        if len(children) >= 3:
            ssubsup = OxmlElement('m:sSubSup')
            e = OxmlElement('m:e')
            for part in _convert_mathml_node(children[0]):
                e.append(part)
            ssubsup.append(e)
            sub = OxmlElement('m:sub')
            for part in _convert_mathml_node(children[1]):
                sub.append(part)
            ssubsup.append(sub)
            sup = OxmlElement('m:sup')
            for part in _convert_mathml_node(children[2]):
                sup.append(part)
            ssubsup.append(sup)
            results.append(ssubsup)
    elif tag == 'mfrac':
        children = list(node)
        if len(children) >= 2:
            f = OxmlElement('m:f')
            num = OxmlElement('m:num')
            for part in _convert_mathml_node(children[0]):
                num.append(part)
            f.append(num)
            den = OxmlElement('m:den')
            for part in _convert_mathml_node(children[1]):
                den.append(part)
            f.append(den)
            results.append(f)
    elif tag == 'msqrt':
        rad = OxmlElement('m:rad')
        deg = OxmlElement('m:deg')
        rad.append(deg)
        e = OxmlElement('m:e')
        for child in node:
            for part in _convert_mathml_node(child):
                e.append(part)
        rad.append(e)
        results.append(rad)
    elif tag == 'mover':
        children = list(node)
        # Treat as base with accent above — simplify to just base + accent text
        for child in children:
            results.extend(_convert_mathml_node(child))
    elif tag == 'munder':
        children = list(node)
        for child in children:
            results.extend(_convert_mathml_node(child))
    elif tag == 'mtable':
        # Flatten table cells into a sequence
        for child in node.iter():
            cname = etree.QName(child.tag).localname if '}' in child.tag else child.tag
            if cname in ('mi', 'mn', 'mo', 'mtext') and child.text:
                results.append(_make_omml_run(child.text))
    else:
        # Fallback: recurse into children
        for child in node:
            results.extend(_convert_mathml_node(child))
        if node.text and not list(node):
            results.append(_make_omml_run(node.text))

    return results

def latex_to_omml(latex_str):
    """Convert LaTeX string to an OMML oMath element"""
    mathml_str = latex2mathml.converter.convert(latex_str)
    root = etree.fromstring(mathml_str.encode())
    omath = OxmlElement('m:oMath')
    for part in _convert_mathml_node(root):
        omath.append(part)
    return omath

def add_display_math(latex_str):
    """Add a display math paragraph (centered formula)"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.line_spacing = 1.5
    try:
        omath_para = OxmlElement('m:oMathPara')
        omath = latex_to_omml(latex_str)
        omath_para.append(omath)
        para._element.append(omath_para)
    except Exception as e:
        # Fallback: render as italic text
        run = para.add_run(latex_str)
        run.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(13)
        print(f'  Math fallback for: {latex_str[:50]}... ({e})')
    return para

def add_inline_math(para, latex_str):
    """Add inline math to an existing paragraph"""
    try:
        omath = latex_to_omml(latex_str)
        para._element.append(omath)
    except Exception:
        run = para.add_run(latex_str)
        run.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(13)

# Helper: add table
def add_table(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
        for para in t.rows[0].cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row):
            t.rows[r_idx+1].cells[c_idx].text = str(cell)
            for para in t.rows[r_idx+1].cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    return t

# Read the markdown
with open('docs/report_vi.md', 'r', encoding='utf-8') as f:
    md = f.read()

# Split into sections by ## headers
sections = {}
current_key = 'header'
current_content = []
for line in md.split('\n'):
    if line.startswith('## '):
        if current_content:
            sections[current_key] = '\n'.join(current_content)
        current_key = line[3:].strip()
        current_content = []
    else:
        current_content.append(line)
if current_content:
    sections[current_key] = '\n'.join(current_content)

print(f"Found {len(sections)} sections:")
for key in sections:
    print(f"  - {key[:60]} ({len(sections[key])} chars)")

# Helper: flush pending table rows
def flush_table(table_rows):
    if len(table_rows) < 2:
        for row in table_rows:
            p(' | '.join(row))
        return
    headers = table_rows[0]
    data = table_rows[1:]
    ncols = len(headers)
    normalized = []
    for row in data:
        while len(row) < ncols:
            row.append('')
        normalized.append(row[:ncols])
    add_table(headers, normalized)

# Helper: render rich text (bold, italic, code) into a paragraph
def _render_rich_text(para, text):
    """Parse **bold**, *italic*, `code` in text and add as runs to para"""
    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
            run = para.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = para.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            continue  # skip default font setting
        else:
            run = para.add_run(token)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)

# Helper: render markdown content block into doc
def render_md(content):
    in_code = False
    code_buf = []
    in_table = False
    table_rows = []

    for line in content.split('\n'):
        # Code blocks
        if line.strip().startswith('```'):
            if in_code:
                add_code('\n'.join(code_buf))
                code_buf = []
                in_code = False
            else:
                if in_table and table_rows:
                    flush_table(table_rows)
                    table_rows = []
                    in_table = False
                in_code = True
            continue
        if in_code:
            code_buf.append(line)
            continue

        # Tables
        if line.strip().startswith('|'):
            cells = [c.strip() for c in line.strip().strip('|').split('|')]
            if all(set(c) <= set('-: ') for c in cells):
                continue
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table and table_rows:
            flush_table(table_rows)
            table_rows = []
            in_table = False

        # Empty lines
        if not line.strip():
            continue

        # Sub-headings
        if line.startswith('#### '):
            h(line[5:].strip(), level=3)
        elif line.startswith('### '):
            h(line[4:].strip(), level=2)
        # Bold lines
        elif line.startswith('**') and line.rstrip().endswith('**'):
            pb(line.strip().strip('*'))
        # Bullet points
        elif line.startswith('- ') or line.startswith('* '):
            bullet = doc.add_paragraph(style='List Bullet')
            _render_rich_text(bullet, line[2:])
        # Numbered items
        elif re.match(r'^\d+\.\s', line):
            item = doc.add_paragraph(style='List Number')
            _render_rich_text(item, line)
        # Display math ($$...$$)
        elif line.strip().startswith('$$'):
            formula = line.strip().strip('$').strip()
            add_display_math(formula)
        # Normal text (with possible inline math $...$ and **bold**)
        else:
            # Build a rich paragraph with bold, italic, code, and inline math
            para = doc.add_paragraph()
            para.paragraph_format.first_line_indent = Cm(1.27)
            para.paragraph_format.line_spacing = 1.5

            # Split by inline math first, then handle bold/italic/code within text parts
            if '$' in line:
                parts = re.split(r'(?<!\$)\$(?!\$)(.+?)(?<!\$)\$(?!\$)', line)
            else:
                parts = [line]

            for i, part in enumerate(parts):
                if '$' in line and i % 2 == 1:
                    # Inline math
                    add_inline_math(para, part)
                else:
                    # Rich text: split by **bold**, *italic*, `code`
                    tokens = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', part)
                    for token in tokens:
                        if not token:
                            continue
                        if token.startswith('**') and token.endswith('**'):
                            run = para.add_run(token[2:-2])
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                        elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
                            run = para.add_run(token[1:-1])
                            run.italic = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)
                        elif token.startswith('`') and token.endswith('`'):
                            run = para.add_run(token[1:-1])
                            run.font.name = 'Courier New'
                            run.font.size = Pt(11)
                        else:
                            run = para.add_run(token)
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(13)

    # Flush remaining
    if in_code and code_buf:
        add_code('\n'.join(code_buf))
    if in_table and table_rows:
        flush_table(table_rows)


# ============================================================
# BUILD DOCUMENT
# ============================================================

# --- LỜI CẢM ƠN ---
h('LỜI CẢM ƠN', 0)
p('Nhóm em xin được gửi lời cảm ơn chân thành đến các thầy cô trong Khoa đã truyền đạt cho nhóm em những kiến thức quý báu trong quá trình học tập tại trường. Đặc biệt, em xin gửi lời cảm ơn đến thầy giáo hướng dẫn — người đã truyền đạt những kiến thức quý báu làm nền tảng để nhóm em xây dựng và hoàn thiện báo cáo.')
p('Mặc dù có sự quan tâm, giúp đỡ của các thầy nhưng do kiến thức và kinh nghiệm của nhóm còn hạn chế nên Báo cáo chuyên đề của nhóm không thể tránh khỏi những thiếu sót. Chúng em rất mong nhận được những ý kiến đóng góp từ quý thầy cô để báo cáo được hoàn thiện hơn.')
p('Chúng em xin chân thành cảm ơn!')

doc.add_page_break()

# --- MỞ ĐẦU ---
h('MỞ ĐẦU', 0)

h('1. Lý do lựa chọn đề tài', 1)
p('Trong bối cảnh an ninh mạng ngày càng phức tạp, tấn công lừa đảo trực tuyến (phishing) và mã độc sử dụng tên miền sinh tự động (DGA — Domain Generation Algorithm) đã trở thành hai trong những mối đe dọa nghiêm trọng nhất. Theo báo cáo của Anti-Phishing Working Group (APWG), năm 2023 ghi nhận hơn 4.7 triệu cuộc tấn công phishing, mức cao nhất từng được ghi nhận. FBI báo cáo thiệt hại từ phishing vượt 10 tỷ USD trong năm 2022.')
p('Các phương pháp truyền thống như danh sách đen (blacklist) không thể cập nhật kịp thời với các URL mới — một trang phishing có thể tồn tại chỉ vài giờ trước khi bị gỡ xuống. Tương tự, các dòng malware như Conficker, CryptoLocker sử dụng DGA để tạo ra hàng chục nghìn tên miền mỗi ngày, khiến việc chặn thủ công trở nên bất khả thi.')
p('Xuất phát từ thực tế này, việc nghiên cứu và ứng dụng Machine Learning để phát hiện URL phishing và domain DGA trong thời gian thực là cần thiết và cấp bách. Đây là lý do nhóm em lựa chọn đề tài "Nghiên cứu, ứng dụng Machine Learning để xây dựng hệ thống phát hiện Phishing URL và DGA Domain, tích hợp Chrome Extension cảnh báo người dùng".')

h('2. Mục tiêu và nhiệm vụ nghiên cứu', 1)
p('Chuyên đề tập trung xây dựng một hệ thống phát hiện phishing URL và DGA domain sử dụng mô hình XGBoost, với các nhiệm vụ cụ thể:')
p('- Nghiên cứu lý thuyết về Machine Learning, đặc biệt là Gradient Boosting và XGBoost.')
p('- Khám phá và phân tích hai bộ dữ liệu: PhiUSIIL Phishing URL Dataset (235,795 mẫu) và ExtraHop DGA Detection Dataset (16.2 triệu mẫu).')
p('- Huấn luyện hai mô hình XGBoost riêng biệt: mô hình phát hiện phishing (25 đặc trưng URL-only) và mô hình phát hiện DGA (71 đặc trưng).')
p('- Xây dựng backend API (FastAPI) kết hợp kết quả hai mô hình.')
p('- Tích hợp vào Chrome Extension để cảnh báo người dùng trực quan khi duyệt web.')
p('- Phân tích và giải thích mô hình bằng SHAP (SHapley Additive exPlanations).')

h('3. Đối tượng và phạm vi nghiên cứu', 1)
p('Đối tượng nghiên cứu: Thuật toán XGBoost và ứng dụng trong phân loại URL phishing và domain DGA. Phương pháp giải thích mô hình SHAP.')
p('Phạm vi nghiên cứu: Sử dụng 25 đặc trưng URL-only (pre-fetch) cho phishing detection — không cần tải trang web. Sử dụng 71 đặc trưng (64 mã hóa ký tự + 7 thống kê) cho DGA detection. Triển khai trên Chrome Extension với kiến trúc hai tầng (client-side heuristic + backend ML).')

h('4. Phương pháp nghiên cứu', 1)
p('Phương pháp lý thuyết: Nghiên cứu tài liệu về Decision Tree, Gradient Boosting, XGBoost, và các chỉ số đánh giá (Accuracy, Precision, Recall, F1, ROC-AUC). Tìm hiểu SHAP values và TreeSHAP.')
p('Phương pháp thực nghiệm: Huấn luyện và đánh giá mô hình trên dữ liệu thực. Xây dựng hệ thống backend + extension. Phân tích SHAP để giải thích kết quả. So sánh Feature Importance (gain) với SHAP values.')

h('5. Ý nghĩa khoa học và thực tiễn', 1)
p('Ý nghĩa khoa học: Chuyên đề cung cấp phân tích chi tiết về XGBoost trong bài toán phát hiện URL độc hại, so sánh giữa Feature Importance và SHAP, và chứng minh rằng chỉ cần đặc trưng từ chuỗi URL (không cần tải trang) đã đủ để đạt accuracy 99.98%.')
p('Ý nghĩa thực tiễn: Hệ thống cho phép phát hiện phishing và DGA trong thời gian thực, bảo vệ người dùng ngay khi họ di chuột qua link — trước khi click. Chrome Extension hoạt động trên mọi trang web, tích hợp cả phân tích nhanh (client-side) và phân tích sâu (backend ML).')

h('6. Bố cục chuyên đề', 1)
p('Ngoài phần mở đầu, kết luận và tài liệu tham khảo, nội dung chuyên đề gồm 4 chương:')
p('Chương 1: Tổng quan nghiên cứu liên quan — Khảo sát tình hình nghiên cứu hiện tại về phát hiện DGA và phishing URL, bao gồm các công trình học thuật, sản phẩm công nghiệp, và so sánh với phương pháp của dự án.')
p('Chương 2: Lý thuyết chung — Trình bày kiến thức nền tảng về Machine Learning, Decision Tree, Gradient Boosting, XGBoost, các chỉ số đánh giá, và phương pháp giải thích mô hình SHAP.')
p('Chương 3: Dữ liệu và thiết kế hệ thống — Phân tích hai bộ dữ liệu, kiến trúc hệ thống tổng thể, thiết kế Chrome Extension và backend API.')
p('Chương 4: Kết quả thực nghiệm — Kết quả huấn luyện hai mô hình, phân tích SHAP, so sánh và đánh giá.')

doc.add_page_break()

# --- CHƯƠNG 1: TỔNG QUAN NGHIÊN CỨU LIÊN QUAN ---
h('CHƯƠNG 1: TỔNG QUAN NGHIÊN CỨU LIÊN QUAN', 0)

for key in sections:
    if '7.' in key and 'Tổng quan' in key:
        print(f"Rendering Chapter 1 (Literature Review) from: {key}")
        render_md(sections[key])
        break

doc.add_page_break()

# --- CHƯƠNG 2: LÝ THUYẾT CHUNG ---
h('CHƯƠNG 2: LÝ THUYẾT CHUNG', 0)

# Find and render Section 3
for key in sections:
    if '3.' in key and 'Kiến thức' in key:
        print(f"Rendering Chapter 2 from: {key}")
        render_md(sections[key])
        break

doc.add_page_break()

# --- CHƯƠNG 3: DỮ LIỆU VÀ THIẾT KẾ HỆ THỐNG ---
h('CHƯƠNG 3: DỮ LIỆU VÀ THIẾT KẾ HỆ THỐNG', 0)

# Section 2: Dataset exploration
for key in sections:
    if '2.' in key and ('Khám phá' in key or 'Dataset' in key):
        print(f"Rendering Chapter 2 datasets from: {key}")
        render_md(sections[key])
        break

# Section 1.5: System architecture
for key in sections:
    if '1.' in key and 'Đặt vấn đề' in key:
        content = sections[key]
        if '### 1.5' in content:
            arch_start = content.index('### 1.5')
            print("Rendering system architecture")
            render_md(content[arch_start:])
        break

# Section 5: Chrome Extension architecture (overview only, no code details)
for key in sections:
    if '5.' in key and ('Kiến trúc' in key or 'Extension' in key):
        content = sections[key]
        # Only render 5.1 (overview) and 5.4 (execution flow), skip 5.2/5.3/5.5 code details
        # Find section boundaries
        parts_to_render = []
        # 5.1 Tổng quan
        if '### 5.1' in content:
            start = content.index('### 5.1')
            end = content.index('### 5.2') if '### 5.2' in content else len(content)
            parts_to_render.append(content[start:end])
        # 5.4 Luồng thực thi
        if '### 5.4' in content:
            start = content.index('### 5.4')
            end = content.index('### 5.5') if '### 5.5' in content else len(content)
            parts_to_render.append(content[start:end])
        # 5.6 Giao diện trực quan
        if '### 5.6' in content:
            start = content.index('### 5.6')
            parts_to_render.append(content[start:])
        print(f"Rendering Chrome Extension (overview only): {len(parts_to_render)} subsections")
        for part in parts_to_render:
            render_md(part)
        break

# Section 4.4: How models combine
for key in sections:
    if '4.' in key and ('Kết quả' in key or 'đánh giá' in key.lower()):
        content = sections[key]
        if '### 4.4' in content:
            combine_start = content.index('### 4.4')
            combine_end = content.index('### 4.5') if '### 4.5' in content else len(content)
            print("Rendering model combination")
            render_md(content[combine_start:combine_end])
        break

doc.add_page_break()

# --- CHƯƠNG 4: KẾT QUẢ THỰC NGHIỆM ---
h('CHƯƠNG 4: KẾT QUẢ THỰC NGHIỆM', 0)

for key in sections:
    if '4.' in key and ('Kết quả' in key or 'đánh giá' in key.lower()):
        content = sections[key]
        # Render 4.1, 4.2, 4.3
        if '### 4.4' in content:
            print("Rendering training results (4.1-4.3)")
            render_md(content[:content.index('### 4.4')])
        else:
            render_md(content)

        break

doc.add_page_break()

# --- KẾT LUẬN ---
h('KẾT LUẬN', 0)
p('Chuyên đề đã hoàn thành việc nghiên cứu, xây dựng và đánh giá hệ thống phát hiện Phishing URL và DGA domain sử dụng XGBoost, với các kết quả chính:')
p('Thứ nhất, về mặt lý thuyết, chuyên đề đã trình bày chi tiết kiến thức về Decision Tree, Gradient Boosting và XGBoost với các ví dụ số cụ thể, giúp người đọc hiểu rõ cơ chế hoạt động từ cơ bản đến nâng cao. Phương pháp giải thích mô hình SHAP cũng được giới thiệu và áp dụng.')
p('Thứ hai, về mặt thực nghiệm, mô hình phishing đạt accuracy 99.98% (Recall 100%) chỉ với 25 đặc trưng URL-only, không cần tải trang web. Mô hình DGA đạt accuracy 97.45%, vượt mốc 94.8% của ExtraHop Networks. Phân tích SHAP tiết lộ rằng URLSimilarityIndex là đặc trưng quan trọng nhất cho phishing (43.7% SHAP), và consonants_consec là đặc trưng thống kê quan trọng nhất cho DGA (48.2% SHAP trong nhóm).')
p('Thứ ba, về mặt kỹ thuật, hệ thống Chrome Extension với kiến trúc hai tầng (Fast Detector phía client + Backend XGBoost) cho phép phát hiện URL độc hại trong thời gian thực, cảnh báo người dùng trực quan trước khi họ click vào link.')
p('Hạn chế và hướng phát triển: Mô hình hiện tại chỉ phân tích đặc trưng từ chuỗi URL/domain, chưa xem xét nội dung trang web hoặc ngữ cảnh. Trong tương lai, có thể tích hợp thêm đặc trưng post-fetch, sử dụng deep learning (LSTM, Transformer) cho DGA detection, và mở rộng hệ thống hỗ trợ nhiều trình duyệt.')

doc.add_page_break()

# --- TÀI LIỆU THAM KHẢO ---
h('TÀI LIỆU THAM KHẢO', 0)
for key in sections:
    if '6.' in key and 'Tài liệu' in key:
        print(f"Rendering references from: {key}")
        render_md(sections[key])
        break

# Save
doc.save('docs/report_vi.docx')
print('\nDone! Saved to docs/report_vi.docx')
